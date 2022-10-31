import sys

from docx import Document


def invitation(place1, time1, name1):
    document = Document()
    event = 'Концерт \'International Women\'s Day\''
    for i in range(len(name1)):
        document.add_heading(f'Добрый день, {name1[i]}', 0)

        p = document.add_paragraph(
            f'Поздравляем Вас с 8 мартом и приглашаем на {event}, которое состоится в {place1}.')
        r = document.add_paragraph(f'Начало в {time1}.')
        if i != len(name1) - 1:
            document.add_page_break()

    document.save('test.docx')


place = input().lower()  # Пишите, пожалуйста, без предлога 'в'
time = input().lower()  # Пишите, пожалуйста, без предлога 'в'
name = list(map(str.strip, sys.stdin))
invitation(place, time, name)
