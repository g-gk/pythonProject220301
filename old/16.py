from docx import Document
from docx.shared import Inches
from sys import stdin

place = input()
time = input()
document = Document()  # создаем документ
for name in stdin:
    document.add_paragraph(place)
    document.add_paragraph(time)
    document.add_paragraph(name)
    document.add_page_break()

document.save('test.docx')
