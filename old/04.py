import docx

from docx.shared import Inches
from sys import stdin

document = docx.Document()

a = input()
b = input()

while True:
    g = input()
    if g == '':
        break
    document.add_heading('С 8 марта', 0)
    document.add_paragraph('Место: ' + a)
    document.add_paragraph('Время: ' + b)
    p = document.add_paragraph(g)

document.save('test.docx')
