import sys

from docx import Document

document = Document()
where = input().lower()
when = input().lower()
who = sys.stdin.read().split()
names = who[1::3]
for name in names:
    document.add_heading(f'Приглашение для {name}')
    document.add_paragraph(f'Дорогая {name}, приглашаем вас на классное мероприятие к 8 марта.\n'
                           f'Оно пройдёт {when} {where}.\n'
                           f'Очень надеемся, что вы его посетите.')
    document.add_page_break()
document.save('invitations.docx')
