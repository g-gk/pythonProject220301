from docx import Document
import sys

d = list(map(str.strip, sys.stdin))
print(d)
document = Document()
for x in d[2:]:
    document.add_heading(x + '!', 0)
    document.add_paragraph(
        'Приглашаем тебя на мероприятие к празднику 8 марта в ' + str(d[1])
        + ' к ' + str(d[0]) + '!')
    if x != d[-1]:
        document.add_page_break()

document.save('test.docx')
