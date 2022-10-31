from docx import Document
import sys

# вводим информацию о месте и дате проведения
place = input()
time = input()
# с помощью ввода sys.stdin все имена добавляем в список name
name = []
for line in sys.stdin:
    if line == '\n':
        break
    c = line.rstrip('\n')
    name.append(c)
# создаем документ
doc = Document()
# для каждого имени заполняем приглашение
for i in range(len(name)):
    doc.add_heading(f'Добрый день, {name[i]}!', 0)
    doc.add_paragraph(('Приглашаю тебя на мероприятие в честь 8 марта '
                       f'которое состоится {place} {time}.'))
    if i != len(name) - 1:
        doc.add_page_break()
doc.save("res.docx")
