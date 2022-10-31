import docx as doc
import sys

data = list(map(str.strip, sys.stdin))
place = data[0]
time = data[1]
names = data[2:]
for i in sys.stdin:
    names.append(i.rstrip("\n"))
document = doc.Document()
for i in names:
    document.add_heading("ПРИГЛАШЕНИЕ", 0)
    document.add_heading("Дорогая " + i + ", приглашаем тебя", level=1)
    p = document.add_paragraph("На мероприятие ")
    p = document.add_paragraph(place)
    p = document.add_paragraph(time)
    document.add_paragraph("С 8 Марта!")
    document.add_page_break()
document.save("document.docx")
